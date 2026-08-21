import os
import uuid
from pathlib import Path
from typing import Callable, Optional

import pdfplumber
from pypdf import PdfReader
import docx
import pandas as pd

from core.models import DocumentIngere
from core.paths import app_data_dir

SUPPORTED_EXT = {".pdf", ".docx", ".xlsx", ".csv"}
OCR_THRESHOLD_CHARS = 50


def _extract_pdf_text(path: Path) -> tuple[str, bool]:
    text_parts: list[str] = []
    ocr_used = False
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if len(page_text.strip()) < OCR_THRESHOLD_CHARS:
                    ocr_text = _ocr_page(path, page.page_number - 1)
                    if ocr_text:
                        page_text = ocr_text
                        ocr_used = True
                text_parts.append(page_text)
    except Exception:
        try:
            reader = PdfReader(str(path))
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if len(page_text.strip()) < OCR_THRESHOLD_CHARS:
                    ocr_text = _ocr_page(path, i)
                    if ocr_text:
                        page_text = ocr_text
                        ocr_used = True
                text_parts.append(page_text)
        except Exception as exc:
            return f"[ERREUR EXTRACTION PDF: {exc}]", False
    return "\n".join(text_parts), ocr_used


def _ocr_page(path: Path, page_index: int) -> Optional[str]:
    try:
        import pytesseract
        from pdf2image import convert_from_path

        images = convert_from_path(str(path), first_page=page_index + 1, last_page=page_index + 1)
        if not images:
            return None
        return pytesseract.image_to_string(images[0], lang="fra")
    except Exception:
        return None


def _extract_docx_text(path: Path) -> str:
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx_text(path: Path) -> str:
    try:
        sheets = pd.read_excel(str(path), sheet_name=None, header=None, engine="openpyxl")
        parts = []
        for name, df in sheets.items():
            parts.append(f"--- Feuille: {name} ---")
            parts.append(df.to_csv(index=False, header=False))
        return "\n".join(parts)
    except Exception as exc:
        return f"[ERREUR EXTRACTION XLSX: {exc}]"


def _extract_csv_text(path: Path) -> str:
    try:
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                df = pd.read_csv(str(path), sep=None, engine="python", encoding=enc)
                return df.to_csv(index=False)
            except Exception:
                continue
        return path.read_text(errors="ignore")
    except Exception as exc:
        return f"[ERREUR EXTRACTION CSV: {exc}]"


def extract_document(path: Path) -> DocumentIngere:
    ext = path.suffix.lower()
    doc_id = f"{path.stem}_{uuid.uuid4().hex[:8]}"
    ocr_used = False
    erreur = None
    texte = ""
    try:
        if ext == ".pdf":
            texte, ocr_used = _extract_pdf_text(path)
        elif ext == ".docx":
            texte = _extract_docx_text(path)
        elif ext == ".xlsx":
            texte = _extract_xlsx_text(path)
        elif ext == ".csv":
            texte = _extract_csv_text(path)
        else:
            erreur = f"Extension non supportee: {ext}"
    except Exception as exc:
        erreur = str(exc)

    doc = DocumentIngere(
        doc_id=doc_id,
        chemin=str(path),
        type_fichier=ext.lstrip("."),
        texte_brut=texte,
        nb_caracteres=len(texte),
        ocr_utilise=ocr_used,
        erreur=erreur,
    )
    _save_raw_text(doc)
    return doc


def _save_raw_text(doc: DocumentIngere) -> None:
    staging = app_data_dir() / "staging" / "raw_texts"
    out_path = staging / f"{doc.doc_id}.txt"
    out_path.write_text(doc.texte_brut, encoding="utf-8", errors="ignore")


def scan_folder(
    folder: Path,
    progress_callback: Optional[Callable[[int, int, str], None]] = None,
    cancel_flag: Optional[Callable[[], bool]] = None,
) -> list[DocumentIngere]:
    """Parcours recursif du dossier source. Doit etre appele dans un thread
    d'arriere-plan par l'appelant UI pour ne jamais geler l'interface."""
    files = [
        p for p in folder.rglob("*")
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXT
    ]
    results: list[DocumentIngere] = []
    total = len(files)
    for i, path in enumerate(files, start=1):
        if cancel_flag and cancel_flag():
            break
        if progress_callback:
            progress_callback(i, total, path.name)
        results.append(extract_document(path))
    return results
